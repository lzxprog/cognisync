import logging
import os
import warnings
from io import BytesIO
from pathlib import Path
from typing import Union, Optional
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.pdfparser import PDFSyntaxError
import chardet

from config import MAX_FILE_SIZE

# 获取日志记录器
logger = logging.getLogger(__name__)

# 配置日志
warnings.filterwarnings("ignore", category=UserWarning)  # 禁用PDFMiner的警告


# 辅助函数：日志输出处理
def log_file_info(file_path: str, file_ext: str) -> None:
    """记录文件路径和扩展名信息"""
    logger.info(f"Processing {file_path}")
    logger.info(f"File extension: {file_ext}")


def extract_file_content(file_path: str) -> str:
    """提取文件内容，带格式校验"""
    file_ext = os.path.splitext(file_path)[1].lower()
    log_file_info(file_path, file_ext)

    handlers = {
        '.docx': extract_text_from_docx,
        '.pdf': extract_text_from_pdf,
        '.txt': lambda p: Path(p).read_text(encoding='utf-8')
    }

    if file_ext not in handlers:
        raise ValueError(f"Unsupported file type: {file_ext}")

    content = handlers[file_ext](file_path)
    if not content.strip():
        raise ValueError("Empty file content")
    return content.strip()


def extract_text_from_docx(docx_file: Union[BytesIO, str, Path]) -> str:
    """
    增强版DOCX文本提取，支持表格内容提取和清理
    """
    try:
        # 统一处理路径对象
        if isinstance(docx_file, Path):
            docx_file = str(docx_file)

        # 验证文件大小
        if _check_file_size(docx_file):
            raise ValueError("File size exceeds limit")

        doc = Document(docx_file)
        text_content = []

        # 提取段落文本并清理
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:  # 排除空段落
                text_content.append(para_text)

        # 提取表格内容并清理
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:  # 排除空单元格
                        text_content.append(cell_text)

        # 合并并清理文本
        full_text = "\n".join(text_content).strip()

        # 进一步清理：去除多余的空行
        cleaned_text = "\n".join([line.strip() for line in full_text.splitlines() if line.strip()])

        if not cleaned_text:
            raise ValueError("Extracted text is empty")

        return cleaned_text

    except PackageNotFoundError:
        raise ValueError("Invalid or corrupted DOCX file")
    except Exception as e:
        logger.error(f"DOCX processing failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to process DOCX file: {str(e)}")


def extract_text_from_pdf(pdf_file: Union[BytesIO, str, Path]) -> str:
    """
    增强版PDF文本提取，支持加密检测和备用解析器
    """
    try:
        pdf_file = str(pdf_file) if isinstance(pdf_file, Path) else pdf_file

        # 文件大小检查
        if _check_file_size(pdf_file):
            raise ValueError("File size exceeds limit")

        # 尝试PDFMiner提取
        text = pdfminer_extract_text(pdf_file)
        clean_text = _clean_pdf_text(text)

        if clean_text:
            return clean_text

        # 如果PDFMiner失败，尝试PyMuPDF
        try:
            import fitz  # 延迟导入
            return _extract_with_pymupdf(pdf_file)
        except ImportError:
            logger.warning("PyMuPDF not installed, using fallback method")
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {str(e)}")

        # 最终检查
        if not clean_text:
            if _is_scanned_pdf(pdf_file):
                raise ValueError("Scanned PDF (image-based) not supported")
            raise ValueError("No extractable text found")
        return clean_text

    except PDFSyntaxError:
        raise ValueError("Invalid or corrupted PDF file")
    except Exception as e:
        logger.error(f"PDF processing failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to process PDF file: {str(e)}")


def extract_text_from_txt(file_path: Union[str, Path], encodings: Optional[list] = None) -> str:
    """增强版TXT文件读取，支持自动编码检测"""
    default_encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    encodings = encodings or default_encodings

    try:
        file_path = Path(file_path)
        if _check_file_size(file_path):
            raise ValueError("File size exceeds limit")

        # 首选二进制读取检测编码
        raw_data = file_path.read_bytes()
        result = chardet.detect(raw_data)
        detected_encoding = result['encoding']

        # 尝试检测到的编码优先
        if detected_encoding:
            encodings = [detected_encoding] + encodings

        # 尝试多种编码
        for encoding in encodings:
            try:
                return raw_data.decode(encoding, errors='replace').strip()
            except UnicodeDecodeError:
                continue

        # 所有编码尝试失败
        raise ValueError(f"Failed to decode with tried encodings: {encodings}")

    except Exception as e:
        logger.error(f"TXT processing failed: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to read TXT file: {str(e)}")


def _check_file_size(file_obj: Union[str, Path, BytesIO]) -> bool:
    """验证文件大小"""
    if isinstance(file_obj, (str, Path)):
        file_size = os.path.getsize(file_obj)
    elif isinstance(file_obj, BytesIO):
        file_size = len(file_obj.getvalue())
    else:
        return False
    return file_size > MAX_FILE_SIZE


def _clean_pdf_text(text: str) -> str:
    """清理PDF提取文本"""
    # 移除连续换行和空白字符
    text = "\n".join([line.strip() for line in text.splitlines() if line.strip()])
    # 合并段落
    return text.replace('\n', ' ').replace('  ', ' ').strip()


def _is_scanned_pdf(pdf_path: str) -> bool:
    """简单判断是否为扫描版PDF"""
    try:
        text = pdfminer_extract_text(pdf_path)
        return len(text.strip()) < 50  # 假设可提取文本少于50字符视为扫描件
    except:
        return True


def _extract_with_pymupdf(pdf_file: Union[str, BytesIO]) -> str:
    """使用PyMuPDF作为备用提取器"""
    import fitz  # 确保已安装PyMuPDF

    text = []
    try:
        doc = fitz.open(pdf_file)
        for page in doc:
            text.append(page.get_text("text"))
        return "\n".join(text).strip()
    except fitz.FileDataError:
        raise ValueError("Invalid or encrypted PDF file")
    except Exception as e:
        raise RuntimeError(f"PyMuPDF extraction failed: {str(e)}")